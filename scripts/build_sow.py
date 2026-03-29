"""
SiBoNi CXO Cockpit — Statement of Work (SOW) Generator
Run: python scripts/build_sow.py
Output: SiBoNi_SOW_v1.0.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colour palette ──────────────────────────────────────────────────────────
NAVY    = RGBColor(0x0F, 0x17, 0x2A)
BLUE    = RGBColor(0x1D, 0x4E, 0xD8)
BLUE_LT = RGBColor(0xDB, 0xEA, 0xFE)
VIOLET  = RGBColor(0x6D, 0x28, 0xD9)
AMBER   = RGBColor(0xD9, 0x77, 0x06)
AMBER_LT= RGBColor(0xFF, 0xF8, 0xE7)
GREEN   = RGBColor(0x05, 0x96, 0x69)
GREEN_LT= RGBColor(0xD1, 0xFA, 0xE5)
RED     = RGBColor(0xDC, 0x26, 0x26)
GREY    = RGBColor(0xF1, 0xF5, 0xF9)
GREY_TX = RGBColor(0x47, 0x55, 0x69)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
BLACK   = RGBColor(0x0F, 0x17, 0x2A)

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.top_margin    = Inches(0.9)
section.bottom_margin = Inches(0.9)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)

# ── Styles helper functions ──────────────────────────────────────────────────
def set_run_font(run, size=11, bold=False, italic=False, color=BLACK, name='Calibri'):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color

def set_para_spacing(para, before=4, after=4, line_spacing=None):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)
    if line_spacing:
        para.paragraph_format.line_spacing = Pt(line_spacing)

def add_heading(doc, text, level=1, color=NAVY):
    p = doc.add_paragraph()
    set_para_spacing(p, before=14 if level==1 else 10, after=4)
    run = p.add_run(text)
    sizes = {1:18, 2:14, 3:12, 4:11}
    set_run_font(run, size=sizes.get(level,11), bold=True, color=color)
    return p

def add_para(doc, text, size=10.5, color=BLACK, bold=False, italic=False,
             before=3, after=3, indent=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    set_para_spacing(p, before=before, after=after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def add_bullet(doc, text, size=10.5, indent=0.25):
    p = doc.add_paragraph(style='List Bullet')
    set_para_spacing(p, before=2, after=2)
    p.paragraph_format.left_indent  = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    run = p.add_run(text)
    set_run_font(run, size=size, color=BLACK)
    return p

def add_sub_bullet(doc, text, size=10):
    return add_bullet(doc, text, size=size, indent=0.45)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), kwargs.get(edge, 'single'))
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), kwargs.get(f'{edge}_color', '1D4ED8'))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_rule(doc, color='1D4ED8'):
    p = doc.add_paragraph()
    set_para_spacing(p, before=2, after=2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def cell_para(cell, text, size=10, bold=False, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT,
              italic=False, before=2, after=2):
    if cell.paragraphs and cell.paragraphs[0].text == '':
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.alignment = align
    set_para_spacing(p, before=before, after=after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return p

def add_table_header_row(table, headers, bg='1D4ED8', text_color=WHITE, size=9.5):
    row = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = row.cells[i]
        shade_cell(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, before=4, after=4)
        run = p.add_run(hdr)
        set_run_font(run, size=size, bold=True, color=text_color)

# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════

# Logo / Title block
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p, before=30, after=4)
r = p.add_run("SiBoNi")
set_run_font(r, size=36, bold=True, color=BLUE, name='Calibri')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p, before=0, after=2)
r = p.add_run("CXO COCKPIT  ·  EXECUTIVE INTELLIGENCE PLATFORM")
set_run_font(r, size=10, bold=True, color=GREY_TX)

add_rule(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p, before=18, after=6)
r = p.add_run("STATEMENT OF WORK")
set_run_font(r, size=28, bold=True, color=NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p, before=0, after=4)
r = p.add_run("Software Development & Platform Delivery Agreement")
set_run_font(r, size=13, bold=False, italic=True, color=GREY_TX)

add_rule(doc)

# Metadata table
meta = doc.add_table(rows=8, cols=2)
meta.style = 'Table Grid'
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta.columns[0].width = Inches(2.2)
meta.columns[1].width = Inches(4.3)

meta_data = [
    ("Document Title",    "Statement of Work — SiBoNi CXO Cockpit"),
    ("Document Version",  "v1.1  |  Dated: 29 March 2026"),
    ("Client",            "SiBoNiTech Pvt. Ltd."),
    ("Service Provider",  "Mohit Prasad (Individual / OPC) — with specialist support from"
                          " independent Architects & AWS Engineers"),
    ("Project Name",      "SiBoNi CXO Cockpit — Executive Intelligence Platform"),
    ("Estimated Project Cost", "₹ 90,00,000 (Ninety Lakhs) — Total Estimated"),
    ("SiboniTech Payable","₹ 39,00,000  |  Balance ₹ 51,00,000 = 0% Tech Debt (on Funding)"),
    ("Phase 1 Invoice",   "₹ 9,00,000 — Prototype + MVP  |  Immediate Payout: ₹ 6,00,000"),
]
for i, (k, v) in enumerate(meta_data):
    row = meta.rows[i]
    shade_cell(row.cells[0], 'EFF6FF')
    cell_para(row.cells[0], k, size=9.5, bold=True, color=BLUE)
    cell_para(row.cells[1], v, size=9.5, color=BLACK,
              bold=(i==5 or i==6 or i==7))

set_para_spacing(doc.add_paragraph(), before=20, after=0)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CONFIDENTIAL  ·  FOR AUTHORISED SIGNATORIES ONLY")
set_run_font(r, size=8.5, bold=True, color=GREY_TX)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "TABLE OF CONTENTS", 1, BLUE)
add_rule(doc)

toc_items = [
    ("1.", "Executive Summary", "3"),
    ("2.", "Parties to the Agreement", "3"),
    ("3.", "Project Background & Vision", "4"),
    ("4.", "Scope of Work — Phase 1: Prototype & MVP (Delivered)", "4"),
    ("   4.1", "Technical Architecture Delivered", "4"),
    ("   4.2", "Features & Modules Built", "5"),
    ("   4.3", "Data Connectors & MCP Layer", "6"),
    ("   4.4", "Code Specifications", "6"),
    ("5.", "Full Product Scope — All Phases", "7"),
    ("6.", "Milestone Schedule", "8"),
    ("7.", "Payment Schedule & Plan", "9"),
    ("8.", "AWS Deployment Scope & Cost", "10"),
    ("9.", "Assumptions, Dependencies & Exclusions", "11"),
    ("10.", "Acceptance Criteria", "11"),
    ("11.", "Intellectual Property", "12"),
    ("12.", "Confidentiality", "12"),
    ("13.", "Termination & Dispute Resolution", "12"),
    ("14.", "Signatories", "13"),
]

toc_table = doc.add_table(rows=len(toc_items), cols=3)
toc_table.style = 'Table Grid'
toc_table.columns[0].width = Inches(0.55)
toc_table.columns[1].width = Inches(5.3)
toc_table.columns[2].width = Inches(0.55)

for i, (num, title, page) in enumerate(toc_items):
    row = toc_table.rows[i]
    shade_cell(row.cells[0], 'EFF6FF' if i % 2 == 0 else 'FFFFFF')
    shade_cell(row.cells[1], 'EFF6FF' if i % 2 == 0 else 'FFFFFF')
    shade_cell(row.cells[2], 'EFF6FF' if i % 2 == 0 else 'FFFFFF')
    cell_para(row.cells[0], num, size=9.5, color=BLUE, bold=True if '.' in num.strip() and len(num.strip()) <= 3 else False)
    cell_para(row.cells[1], title, size=9.5, color=BLACK, bold=True if len(num.strip()) <= 3 and num.strip().endswith('.') else False)
    cell_para(row.cells[2], page, size=9.5, color=GREY_TX, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1.  Executive Summary", 1, BLUE)
add_rule(doc)

add_para(doc,
    "This Statement of Work ('SOW') formalises the engagement between SiBoNiTech Pvt. Ltd. "
    "('Client') and Mohit Prasad, an individual technologist and platform architect "
    "('Service Provider'), for the design, development, deployment, and ongoing support of "
    "the SiBoNi CXO Cockpit — an AI-powered Executive Intelligence Platform purpose-built for "
    "C-Suite decision-makers in manufacturing and industrial enterprises. Mohit Prasad works "
    "independently and engages specialist support from freelance Enterprise Architects and "
    "AWS-certified Engineers as required. All responsibilities, IP delivery, and continuity "
    "of the project vest with Mohit Prasad.",
    before=6, after=4)

add_para(doc,
    "As of 29 March 2026, Phase 1 — comprising the Prototype and Minimum Viable Product (MVP) — "
    "has been fully delivered, tested, and committed to the production repository "
    "(github.com/prasad-mohit/executive-dashboard, branch: main). The codebase is live, "
    "buildable, and demonstrable.",
    before=4, after=4)

# Summary box as a 1-cell table
summary_table = doc.add_table(rows=1, cols=1)
summary_table.style = 'Table Grid'
shade_cell(summary_table.rows[0].cells[0], 'EFF6FF')
sc = summary_table.rows[0].cells[0]
cell_para(sc, "KEY FINANCIAL SUMMARY", size=9.5, bold=True, color=BLUE, before=6, after=2)
financials = [
    ("•",  "Total Estimated Project Cost:  ₹ 90,00,000  (Ninety Lakhs)",                          BLACK,  False),
    ("•",  "SiboniTech Total Payable:  ₹ 39,00,000  (Thirty-Nine Lakhs)  — milestone payments",   NAVY,   True),
    ("⚠",  "Tech Debt (0% Interest):  ₹ 51,00,000  — deferred until SiboniTech secures funding",  AMBER,  True),
    ("◆",  "Phase 1 Value (Prototype + MVP):  ₹ 9,00,000  (Nine Lakhs)",                           BLACK,  False),
    ("⚠",  "Immediate / Urgent Payout — Payment 1:  ₹ 6,00,000  (Six Lakhs)   ← DUE NOW",          RED,    True),
    ("•",  "Milestone 1 Balance (Phase 1 sign-off):  ₹ 3,00,000  (Three Lakhs)",                  BLACK,  False),
    ("•",  "Remaining Cash Payments (Phases 2–5):  ₹ 30,00,000  — per milestone schedule",         BLACK,  False),
    ("•",  "Tech Debt Structure:  0% interest, payable upon SiboniTech funding event",              GREY_TX,False),
    ("•",  "Payments to: Mohit Prasad / OPC Company  (bank details furnished separately)",          GREY_TX,False),
]
for icon, line, color, bold in financials:
    p = sc.add_paragraph()
    set_para_spacing(p, before=1, after=1)
    run = p.add_run(icon + "  " + line)
    set_run_font(run, size=9.5, bold=bold, color=color)

set_para_spacing(doc.add_paragraph(), before=4, after=0)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — PARTIES
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2.  Parties to the Agreement", 1, BLUE)
add_rule(doc)

parties = doc.add_table(rows=2, cols=2)
parties.style = 'Table Grid'
parties.columns[0].width = Inches(3.0)
parties.columns[1].width = Inches(3.5)

shade_cell(parties.rows[0].cells[0], '1D4ED8')
shade_cell(parties.rows[0].cells[1], '1D4ED8')
cell_para(parties.rows[0].cells[0], "CLIENT", size=10, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
cell_para(parties.rows[0].cells[1], "SERVICE PROVIDER", size=10, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

r = parties.rows[1]
shade_cell(r.cells[0], 'EFF6FF')
shade_cell(r.cells[1], 'F0FDF4')
c0, c1 = r.cells[0], r.cells[1]
for line in ["SiBoNiTech Pvt. Ltd.", "Registered in India", "Authorised Signatory: [Name]",
             "Designation: [Title]", "Date: _______________", "Signature: _______________"]:
    p = c0.add_paragraph()
    set_para_spacing(p, before=1, after=1)
    run = p.add_run(line)
    bold = line.startswith("SiBoNi")
    set_run_font(run, size=9.5, bold=bold, color=NAVY if bold else BLACK)

for line in ["Mohit Prasad", "Individual / OPC (TBD)", "Platform Architect & Lead Developer",
             "Specialist support: Independent Architects & AWS Engineers",
             "Bank / OPC details furnished separately",
             "Date: _______________", "Signature: _______________"]:
    p = c1.add_paragraph()
    set_para_spacing(p, before=1, after=1)
    run = p.add_run(line)
    bold = line.startswith("Mohit")
    italic = "OPC" in line or "Bank" in line or "Specialist" in line
    set_run_font(run, size=9.5, bold=bold, italic=italic, color=NAVY if bold else (GREY_TX if italic else BLACK))

set_para_spacing(doc.add_paragraph(), before=4, after=0)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 3 — PROJECT BACKGROUND
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3.  Project Background & Vision", 1, BLUE)
add_rule(doc)

add_para(doc,
    "SiBoNi CXO Cockpit is a first-of-its-kind executive intelligence platform designed "
    "specifically for CEOs, COOs, and C-suite leaders in manufacturing, industrial, and "
    "multi-plant enterprises. The platform addresses a universal, high-value problem: the "
    "gap between enterprise data and executive decision-making.",
    before=6, after=4)

add_para(doc, "Core Product Vision:", bold=True, size=10.5, color=NAVY, before=4, after=2)
bullets = [
    "Connect all enterprise systems (ERP, CRM, HR, IoT, Market Feeds) via Model Context Protocol (MCP) connectors",
    "Synthesise signals through always-on AI into executive-grade intelligence — not reports, but decisions",
    "Deliver a CEO morning view: what needs attention today, what to commit, what to hold",
    "Enable analysts to push intelligence directly to the CEO — bypassing email, decks, and reporting cycles",
    "Track executive decisions through to execution with a complete audit trail",
    "Generate 1-click board packs from live, AI-verified data",
]
for b in bullets:
    add_bullet(doc, b)

add_para(doc, "\nTarget Market: Manufacturing & Industrial enterprises, ₹500Cr+ turnover, "
    "multi-plant, multi-region, using SAP ERP + Salesforce CRM + Workday HR.",
    size=10, italic=True, color=GREY_TX, before=6)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 4 — PHASE 1 SCOPE (DELIVERED)
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "4.  Scope of Work — Phase 1: Prototype & MVP (Delivered)", 1, BLUE)
add_rule(doc)

add_para(doc,
    "Phase 1 encompasses the complete Prototype and MVP of the SiBoNi CXO Cockpit. "
    "All items listed in this section have been FULLY DELIVERED and committed to the "
    "production GitHub repository (commit refs: fc7366d → aa95988 → 8ae9aad → 0e50d19 → "
    "2926654, branch: main) as of 29 March 2026.",
    before=6, after=6, color=BLACK)

# 4.1 Technical Architecture
add_heading(doc, "4.1  Technical Architecture Delivered", 2, NAVY)

arch_table = doc.add_table(rows=6, cols=3)
arch_table.style = 'Table Grid'
arch_table.columns[0].width = Inches(1.5)
arch_table.columns[1].width = Inches(2.8)
arch_table.columns[2].width = Inches(2.2)

add_table_header_row(arch_table, ["Layer", "Technology / Framework", "Status"])
arch_rows = [
    ("Frontend SPA",       "React 18 + Vite 5.4.21 + TailwindCSS 3",     "✅ Delivered"),
    ("State Management",   "React Context API (5 contexts) + localStorage", "✅ Delivered"),
    ("Routing",            "React Router v6 — 10 protected routes",        "✅ Delivered"),
    ("Build & Bundle",     "Vite — build time < 5s, bundle optimised",     "✅ Delivered"),
    ("Mock Data Layer",    "CSV, JSON, JS data files — 7 enterprise feeds", "✅ Delivered"),
]
for i, (layer, tech, status) in enumerate(arch_rows):
    row = arch_table.rows[i+1]
    shade_cell(row.cells[0], 'EFF6FF' if i%2==0 else 'FFFFFF')
    shade_cell(row.cells[1], 'EFF6FF' if i%2==0 else 'FFFFFF')
    shade_cell(row.cells[2], 'F0FDF4')
    cell_para(row.cells[0], layer, size=9.5, bold=True, color=BLUE)
    cell_para(row.cells[1], tech, size=9.5)
    cell_para(row.cells[2], status, size=9.5, bold=True, color=GREEN)

set_para_spacing(doc.add_paragraph(), before=6, after=0)

# 4.2 Features & Modules
add_heading(doc, "4.2  Features & Modules Built", 2, NAVY)

modules = [
    ("Authentication & Role-Based Access Control",
     "Salesforce-style login UI. JWT-pattern auth via AuthContext. 4 distinct roles: "
     "executive (CEO), manager (COO), analyst (Senior Data Analyst), admin (Platform Admin). "
     "Role-based route protection via ProtectedRoute component. 7 user accounts configured.",
     "✅"),
    ("Insights Hub — AI-Driven Executive View",
     "KPI selector with 8 metrics (Revenue, EBITDA, On-Time Delivery, Quality, Inventory, "
     "Cash Flow, Workforce, Market Share). Click-to-drill expandable cards for Decisions, "
     "Signals, Risks, and Business Impact. Analyst Intelligence amber banner for CEO role. "
     "Snap & Download functionality. AI narrative per KPI.",
     "✅"),
    ("Decision Hub — COMMIT / HOLD Workflow",
     "Evidence-backed decision cards with ref IDs (DEC-001 to DEC-006). COMMIT and HOLD "
     "actions with persistence via DecisionStateContext. Decision detail page with data evidence, "
     "risk panel, reasoning panel, agent timeline, interoperability feed.",
     "✅"),
    ("Execution Hub — Action Tracking",
     "Committed decision tracking with KPI feedback loop, owner assignment, milestone tracking, "
     "execution status per decision. Connected to DecisionStateContext.",
     "✅"),
    ("Board Brief — 1-Click Board Pack",
     "Auto-generated board brief from live KPIs, committed decisions, top risks, and pinned "
     "analyst insights. Snap-to-export and download functionality.",
     "✅"),
    ("Signals Console — Enterprise Signal Feed",
     "Real-time signal feed across 6 categories (Supply Chain, Revenue, Regulatory, Quality, "
     "Workforce, Market). Severity filtering, confidence scoring, signal drill-down, "
     "source attribution.",
     "✅"),
    ("Admin Panel — Platform Management",
     "3-tab management console: (1) Data Connections — 9 MCP connector cards with Sync Now, "
     "config expand, Add/Remove; (2) System Prompts — 5 pre-loaded prompts with enable/disable, "
     "schedule, Run Now, last output viewer, Add New; (3) Audit Log — all platform actions.",
     "✅"),
    ("Analyst Studio — THE MARKET DIFFERENTIATOR",
     "Natural language prompt builder against all connected live data feeds. 7 domain AI "
     "detectors (Revenue, Supply Chain, Quality, Workforce, Market, Margin, Risk). Structured "
     "insight generation: headline, summary, 4 key metrics, recommendation, confidence bar. "
     "Share with CEO — pushes directly to CEO morning view. Prompt Library with save/recall. "
     "Shared tab with recall control.",
     "✅"),
    ("PromptContext — Intelligence Bridge",
     "localStorage-backed context for full state management: connections[], systemPrompts[], "
     "analystPrompts[], sharedInsights[], auditLog[]. Full CRUD: syncConnection, addSystemPrompt, "
     "toggleSystemPrompt, runSystemPrompt, saveAnalystPrompt, shareToExecutive, "
     "recallSharedInsight, dismissSharedInsight, pinToBoard, markViewed.",
     "✅"),
    ("Sidebar Navigation — Role-Based",
     "Dynamic navigation per role. CEO: 7 items. Analyst: +Analyst Studio (purple, NEW badge). "
     "Admin: +Admin Panel (red). Unviewed analyst insight count badge on Insights Hub nav item.",
     "✅"),
    ("Sales & Technical Deck",
     "Full 11-slide interactive HTML sales deck (siboni-deck.html) — white/light theme, "
     "scroll-based, TOC, progress bar, all E2E flows, competitive comparison, use cases. "
     "Matching 11-slide PowerPoint (siboni-deck.pptx) generated via python-pptx.",
     "✅"),
]

for mod_name, mod_desc, status in modules:
    p = doc.add_paragraph()
    set_para_spacing(p, before=6, after=1)
    r = p.add_run(f"{status}  {mod_name}")
    set_run_font(r, size=10.5, bold=True, color=NAVY)
    add_para(doc, mod_desc, size=9.5, color=GREY_TX, before=1, after=4, indent=0.2)

# 4.3 Data Connectors
doc.add_page_break()
add_heading(doc, "4.3  Data Connectors & MCP Layer (Mock/Simulated in Phase 1)", 2, NAVY)
add_para(doc,
    "The following enterprise data sources have been fully simulated with representative "
    "mock data in Phase 1. All data structures, schemas, KPI computations, and UI bindings "
    "are production-ready. Real API connections are scoped for Phase 2.",
    size=10, color=GREY_TX, before=4, after=6)

conn_table = doc.add_table(rows=10, cols=4)
conn_table.style = 'Table Grid'
conn_table.columns[0].width = Inches(2.2)
conn_table.columns[1].width = Inches(1.8)
conn_table.columns[2].width = Inches(1.5)
conn_table.columns[3].width = Inches(1.0)

add_table_header_row(conn_table, ["Connector / Source", "Data Entities", "Records (Mock)", "Phase 1"])
conn_rows = [
    ("SAP S/4HANA ERP",   "Orders, Inventory, Production, GL, Revenue",     "42,800",  "Simulated ✅"),
    ("Salesforce CRM",    "Accounts, Opportunities, Cases, NPS Scores",      "18,400",  "Simulated ✅"),
    ("Workday HR",        "Headcount, Skills, Performance, Capacity",         "3,200",   "Simulated ✅"),
    ("Reuters Newsfeed",  "Market signals, Supply chain alerts, Regulatory",  "1,240",   "Simulated ✅"),
    ("Bloomberg Markets", "FX rates, Commodity prices, Market indices",       "580",     "Simulated ✅"),
    ("MS Exchange Email", "Email threads, Sentiment, Executive comms",        "89,000",  "Simulated ✅"),
    ("MS Teams",          "Meeting summaries, Action items, Sentiment",       "480",     "Simulated ✅"),
    ("Plant IoT Sensors", "OEE, Machine uptime, Cycle time, Quality",        "3 plants","Beta ⚙"),
    ("Competitor Monitor","Pricing intel, Job postings, Patents, News",       "284",     "Simulated ✅"),
]
for i, (name, entities, records, status) in enumerate(conn_rows):
    row = conn_table.rows[i+1]
    shade_cell(row.cells[0], 'EFF6FF' if i%2==0 else 'FFFFFF')
    shade_cell(row.cells[3], 'F0FDF4')
    cell_para(row.cells[0], name, size=9.5, bold=True, color=BLUE)
    cell_para(row.cells[1], entities, size=9)
    cell_para(row.cells[2], records, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(row.cells[3], status, size=9, bold=True, color=GREEN if "✅" in status else AMBER)

set_para_spacing(doc.add_paragraph(), before=6, after=0)

# 4.4 Code Specifications
add_heading(doc, "4.4  Code Specifications & Deliverables", 2, NAVY)

code_table = doc.add_table(rows=10, cols=3)
code_table.style = 'Table Grid'
code_table.columns[0].width = Inches(1.8)
code_table.columns[1].width = Inches(3.5)
code_table.columns[2].width = Inches(1.2)

add_table_header_row(code_table, ["Metric", "Detail", "Count"])
code_rows = [
    ("Total Source Files",  "React JSX, JS, CSS, config files in src/",                    "40+ files"),
    ("React Components",    "Reusable UI components in src/components/",                    "21 components"),
    ("Page Views",          "Full-page route components in src/pages/",                     "10 pages"),
    ("Context Providers",   "AuthContext, FilterContext, DecisionStateContext, PromptContext, WorkspaceContext", "5 contexts"),
    ("Lines of Code",       "Approximate total LOC across all JSX/JS source files",         "~8,500 LOC"),
    ("Data Files",          "CSV, JSON, JS mock data (ERP, CRM, HR, Signals, Decisions)",   "9 data files"),
    ("Routes",              "Protected SPA routes with role-based access control",           "10 routes"),
    ("Git Commits",         "Versioned commits on main branch with full history",            "5 commits"),
    ("Build Artefacts",     "Vite production build — ✓ built in < 5s, zero errors",        "Build OK ✅"),
]
for i, (metric, detail, count) in enumerate(code_rows):
    row = code_table.rows[i+1]
    shade_cell(row.cells[0], 'EFF6FF' if i%2==0 else 'FFFFFF')
    shade_cell(row.cells[2], 'F0FDF4' if i%2==0 else 'FFFFFF')
    cell_para(row.cells[0], metric, size=9.5, bold=True, color=NAVY)
    cell_para(row.cells[1], detail, size=9.5)
    cell_para(row.cells[2], count, size=9.5, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 5 — FULL PRODUCT SCOPE
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "5.  Full Product Scope — All Phases", 1, BLUE)
add_rule(doc)

phases_overview = [
    ("Phase 1", "Prototype + MVP",
     "9 Lakhs", "Delivered ✅",
     [
         "Complete React SPA with all 10 pages and 21 components",
         "4-role authentication system with full RBAC",
         "8-KPI Insights Hub with AI narratives and drill-down",
         "Decision Hub with COMMIT/HOLD workflow and evidence bundles",
         "Admin Panel: MCP connector management, system prompts, audit log",
         "Analyst Studio: NL prompt → AI insight → Share with CEO",
         "Board Brief auto-generation, Snap & Download",
         "Signals Console with severity scoring and filtering",
         "9 simulated MCP data connectors",
         "Sales deck (HTML + PPTX) — 11 slides, light executive theme",
     ]),
    ("Phase 2", "Real MCP Connector Integration",
     "25 Lakhs", "Upcoming",
     [
         "Live SAP S/4HANA ERP API integration via MCP connector",
         "Live Salesforce CRM OAuth 2.0 integration",
         "Live Workday HR API integration",
         "Reuters and Bloomberg real-time feed subscriptions",
         "Microsoft Exchange and Teams Graph API integration",
         "Plant IoT MQTT/REST gateway integration",
         "Real-time data pipeline with 1–15 minute refresh cycles",
         "MCP server (Node.js) with authentication and rate limiting",
         "Error handling, retry logic, and connection monitoring",
         "Admin Panel connected to live connector status",
     ]),
    ("Phase 3", "Production AI Engine & LLM Integration",
     "22 Lakhs", "Upcoming",
     [
         "Azure OpenAI / OpenAI GPT-4 API integration for insight generation",
         "Replace mock AI with live LLM-powered analysis engine",
         "Context-aware prompt engineering with enterprise data injection",
         "Confidence scoring and source citation from live data",
         "Automated system prompt scheduling with LLM execution",
         "Semantic signal scoring and priority ranking via AI",
         "Board Brief AI narrative generation from live data",
         "Multi-model fallback strategy (GPT-4 → Claude → Llama)",
         "Token cost monitoring and prompt optimisation",
         "AI response caching and performance optimisation",
     ]),
    ("Phase 4", "AWS Cloud Deployment & DevOps",
     "19 Lakhs", "Upcoming",
     [
         "AWS infrastructure provisioning (VPC, ECS Fargate, RDS, ElastiCache)",
         "React SPA deployed to S3 + CloudFront CDN (global edge)",
         "Node.js MCP server on ECS Fargate with auto-scaling",
         "RDS PostgreSQL for decision history, audit logs, user data",
         "ElastiCache Redis for real-time data caching",
         "API Gateway + Lambda for serverless connector endpoints",
         "Route 53 custom domain + ACM SSL certificate",
         "CloudWatch monitoring, alerting, and dashboards",
         "CI/CD pipeline: GitHub Actions → ECR → ECS blue/green deploy",
         "AWS WAF + Shield for security, IAM roles with least privilege",
         "Disaster recovery: Multi-AZ RDS, S3 cross-region replication",
         "Estimated monthly AWS run cost: ₹ 45,000 – ₹ 75,000",
     ]),
    ("Phase 5", "UAT, Go-Live, Training & Hypercare Support",
     "15 Lakhs", "Upcoming",
     [
         "User Acceptance Testing (UAT) with CEO, COO, Analyst, Admin personas",
         "Performance testing: load, stress, and endurance testing",
         "Security penetration testing and VAPT report",
         "SOC 2 Type II compliance documentation",
         "ISO 27001 alignment review",
         "End-user training: CEO persona (1 day), Analyst (2 days), Admin (1 day)",
         "Admin and operations runbook documentation",
         "90-day hypercare support post go-live",
         "SLA agreement: P1 < 4 hrs, P2 < 8 hrs, P3 < 24 hrs",
         "Handover: source code, infrastructure-as-code, data dictionary",
     ]),
]

for phase_id, phase_name, value, status, deliverables in phases_overview:
    # Phase header row
    phase_tbl = doc.add_table(rows=1, cols=1)
    phase_tbl.style = 'Table Grid'
    pc = phase_tbl.rows[0].cells[0]
    bg = 'EFF6FF' if status == "Delivered ✅" else 'FEF3C7'
    shade_cell(pc, bg)
    hdr_p = pc.paragraphs[0]
    set_para_spacing(hdr_p, before=5, after=2)
    r1 = hdr_p.add_run(f"{phase_id}  —  {phase_name}")
    set_run_font(r1, size=11, bold=True, color=BLUE if status == "Delivered ✅" else AMBER)
    r2 = hdr_p.add_run(f"      |      Value: {value}      |      Status: {status}")
    set_run_font(r2, size=9.5, bold=False, color=GREY_TX)

    for d in deliverables:
        add_bullet(doc, d, size=9.5, indent=0.25)
    set_para_spacing(doc.add_paragraph(), before=4, after=0)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 6 — MILESTONE SCHEDULE
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "6.  Milestone Schedule", 1, BLUE)
add_rule(doc)

add_para(doc,
    "The following milestones define deliverables, timelines, and acceptance gates. "
    "Phase 1 milestones are COMPLETED. Phases 2–5 timelines commence upon execution "
    "of this SOW and receipt of Milestone 1 payment.",
    before=6, after=8)

ms_table = doc.add_table(rows=12, cols=5)
ms_table.style = 'Table Grid'
ms_table.columns[0].width = Inches(0.5)
ms_table.columns[1].width = Inches(2.3)
ms_table.columns[2].width = Inches(1.8)
ms_table.columns[3].width = Inches(1.0)
ms_table.columns[4].width = Inches(0.9)

add_table_header_row(ms_table, ["#", "Milestone / Deliverable", "Timeline", "Value (₹)", "Status"])

milestones = [
    ("M1", "Phase 1 — Prototype & MVP\n"
           "• Full SPA (React/Vite) — 10 pages, 21 components\n"
           "• 4-role auth + RBAC\n"
           "• Insights Hub, Decision Hub, Execution Hub\n"
           "• Board Brief, Signals Console, Admin Panel\n"
           "• Analyst Studio (market differentiator)\n"
           "• 9 MCP connector simulations\n"
           "• Sales deck (HTML + PPTX)\n"
           "• GitHub repo — build OK, all commits pushed",
     "COMPLETED\n29 Mar 2026", "6,00,000\n(URGENT\nPAYOUT)", "✅ DONE"),
    ("M2", "Phase 1 — Sign-off & Acceptance\n"
           "• Formal sign-off on delivered Phase 1 scope\n"
           "• Handover of all credentials and access\n"
           "• Phase 1 closeout documentation",
     "Within 7 days\nof M1 payment", "3,00,000", "⏳ Pending"),
    ("M3", "Phase 2A — Core MCP Connectors\n"
           "• SAP ERP + Salesforce CRM live integration\n"
           "• Workday HR + Reuters live feed\n"
           "• MCP server deployed (Node.js)",
     "Weeks 1–4\n(post sign-off)", "8,00,000", "📋 Planned"),
    ("M4", "Phase 2B — Full Connector Suite\n"
           "• Bloomberg + Exchange + Teams + IoT\n"
           "• Real-time pipeline 1–15 min refresh\n"
           "• Admin Panel live connector status",
     "Weeks 5–8", "8,00,000", "📋 Planned"),
    ("M5", "Phase 2C — Competitor Monitor + Delivery\n"
           "• Competitor intel feed live\n"
           "• Full Phase 2 UAT sign-off",
     "Weeks 9–10", "9,00,000", "📋 Planned"),
    ("M6", "Phase 3A — AI Engine Integration\n"
           "• Azure OpenAI / GPT-4 API wired\n"
           "• Live LLM insight generation in Analyst Studio\n"
           "• System prompt scheduler with LLM execution",
     "Weeks 11–14", "11,00,000", "📋 Planned"),
    ("M7", "Phase 3B — AI Engine Full + Optimisation\n"
           "• Board Brief AI narrative from live data\n"
           "• Confidence scoring + source citation live\n"
           "• Token cost monitoring + prompt optimisation",
     "Weeks 15–18", "11,00,000", "📋 Planned"),
    ("M8", "Phase 4A — AWS Infrastructure Setup\n"
           "• VPC, ECS Fargate, RDS, ElastiCache provisioned\n"
           "• React SPA on S3 + CloudFront\n"
           "• CI/CD pipeline: GitHub → ECR → ECS",
     "Weeks 19–22", "10,00,000", "📋 Planned"),
    ("M9", "Phase 4B — AWS Full Deployment\n"
           "• MCP server deployed to ECS Fargate\n"
           "• API Gateway + Lambda connectors\n"
           "• CloudWatch monitoring + WAF + Shield\n"
           "• Route 53 custom domain + SSL",
     "Weeks 23–26", "9,00,000", "📋 Planned"),
    ("M10","Phase 5A — UAT & Security\n"
           "• Full UAT across 4 personas\n"
           "• VAPT penetration testing\n"
           "• SOC 2 + ISO 27001 documentation",
     "Weeks 27–30", "8,00,000", "📋 Planned"),
    ("M11","Phase 5B — Training & Go-Live\n"
           "• End-user training (CEO, Analyst, Admin)\n"
           "• Go-live sign-off\n"
           "• 90-day hypercare support begins\n"
           "• Full handover: code, infra, documentation",
     "Week 31 onward", "7,00,000", "📋 Planned"),
]

row_bgs = ['EFF6FF','F0FDF4','FFFBEB','FFFBEB','FFFBEB','FDF4FF','FDF4FF','FEF2F2','FEF2F2','F0FDF4','F0FDF4']
for i, (ms, desc, timeline, value, status) in enumerate(milestones):
    row = ms_table.rows[i+1]
    bg = row_bgs[i] if i < len(row_bgs) else 'FFFFFF'
    for c in row.cells:
        shade_cell(c, bg)
    cell_para(row.cells[0], ms, size=9, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    # Multi-line description
    lines = desc.strip().split('\n')
    cell_para(row.cells[1], lines[0], size=9.5, bold=True, color=NAVY)
    for l in lines[1:]:
        p2 = row.cells[1].add_paragraph()
        set_para_spacing(p2, before=0, after=0)
        r2 = p2.add_run(l)
        set_run_font(r2, size=8.5, color=GREY_TX)
    cell_para(row.cells[2], timeline, size=9, color=GREY_TX, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(row.cells[3], value, size=9.5, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    sc = GREEN if "✅" in status else (AMBER if "⏳" in status else GREY_TX)
    cell_para(row.cells[4], status, size=9, bold=True, color=sc, align=WD_ALIGN_PARAGRAPH.CENTER)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 7 — PAYMENT SCHEDULE
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "7.  Payment Schedule & Plan", 1, BLUE)
add_rule(doc)

add_para(doc,
    "All payments are in Indian Rupees (INR). GST at applicable rates will be charged "
    "additionally on each invoice. Payments are due within the specified timelines from "
    "invoice date. A delay of more than 14 days may result in work suspension.",
    before=6, after=8, size=10, color=GREY_TX)

pay_table = doc.add_table(rows=13, cols=5)
pay_table.style = 'Table Grid'
pay_table.columns[0].width = Inches(0.45)
pay_table.columns[1].width = Inches(1.5)
pay_table.columns[2].width = Inches(2.2)
pay_table.columns[3].width = Inches(1.2)
pay_table.columns[4].width = Inches(1.15)

add_table_header_row(pay_table, ["#", "Milestone", "Trigger / Due Date", "Amount (₹)", "Cumulative (₹)"])

payments = [
    ("P1", "M1 — Prototype +\nMVP Delivery",
     "IMMEDIATE — Due on SOW execution\n⚠ URGENT — Transfer to Mohit Prasad / OPC",
     "6,00,000", "6,00,000", True),
    ("P2", "M2 — Phase 1\nSign-off",
     "Within 7 days of M1 payment\nand formal acceptance sign-off",
     "3,00,000", "9,00,000", False),
    ("P3", "M3 — Phase 2A\nDelivery",
     "On delivery of SAP + CRM\nlive connector milestone",
     "8,00,000", "17,00,000", False),
    ("P4", "M4 — Phase 2B\nDelivery",
     "On delivery of full\nconnector suite",
     "8,00,000", "25,00,000", False),
    ("P5", "M5 — Phase 2\nSign-off",
     "On Phase 2 UAT acceptance\n+ final connector milestone",
     "9,00,000", "34,00,000", False),
    ("P6", "M6 — Phase 3A\nAI Engine",
     "On live LLM integration\ndelivery",
     "11,00,000", "45,00,000", False),
    ("P7", "M7 — Phase 3B\nAI Full",
     "On Phase 3 completion\nand sign-off",
     "11,00,000", "56,00,000", False),
    ("P8", "M8 — Phase 4A\nAWS Infra",
     "On AWS infrastructure\nprovisioning complete",
     "10,00,000", "66,00,000", False),
    ("P9", "M9 — Phase 4B\nAWS Deploy",
     "On full production\ndeployment sign-off",
     "9,00,000", "75,00,000", False),
    ("P10","M10 — UAT &\nSecurity",
     "On UAT completion +\nVAPT report delivery",
     "8,00,000", "83,00,000", False),
    ("P11","M11 — Go-Live\n& Hypercare",
     "On go-live sign-off +\ntraining completion",
     "7,00,000", "90,00,000", False),
]

row_bgs_p = ['FEF2F2','F0FDF4'] + ['FFFFFF']*9
for i, row_data in enumerate(payments):
    pn, ms_label, trigger, amount, cumulative, urgent = row_data
    row = pay_table.rows[i+1]
    bg = 'FEF2F2' if urgent else ('F0FDF4' if i==1 else ('FFFFFF' if i%2==0 else 'F8FAFC'))
    for c in row.cells:
        shade_cell(c, bg)
    cell_para(row.cells[0], pn, size=9.5, bold=True, color=RED if urgent else BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(row.cells[1], ms_label, size=9.5, bold=True, color=RED if urgent else NAVY)
    cell_para(row.cells[2], trigger, size=9, color=RED if urgent else GREY_TX, bold=urgent)
    cell_para(row.cells[3], amount, size=10, bold=True, color=RED if urgent else NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(row.cells[4], cumulative, size=9.5, color=GREY_TX, align=WD_ALIGN_PARAGRAPH.CENTER)

# Total row
tot_row = pay_table.rows[12]
for c in tot_row.cells:
    shade_cell(c, '1D4ED8')
cell_para(tot_row.cells[0], "—",                       size=10, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
cell_para(tot_row.cells[1], "TOTAL (Estimated)",        size=10, bold=True, color=WHITE)
cell_para(tot_row.cells[2], "Cash: ₹39L  |  Tech Debt: ₹51L @ 0%", size=9, color=WHITE)
cell_para(tot_row.cells[3], "90,00,000",               size=11, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
cell_para(tot_row.cells[4], "Est. Total",              size=9.5, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

set_para_spacing(doc.add_paragraph(), before=6, after=0)

# Payment terms callout
pt_box = doc.add_table(rows=1, cols=1)
pt_box.style = 'Table Grid'
shade_cell(pt_box.rows[0].cells[0], 'FFF7ED')
ptc = pt_box.rows[0].cells[0]
cell_para(ptc, "PAYMENT TERMS", size=10, bold=True, color=AMBER, before=5, after=3)
pt_items = [
    "Payments by NEFT/RTGS to: Mohit Prasad (Individual) / OPC Company — bank details furnished separately",
    "SiboniTech total cash outflow: ₹ 39,00,000 across all milestones",
    "Remaining ₹ 51,00,000 recorded as 0% interest Tech Debt on SiboniTech's books — payable on funding event",
    "GST (18%) applicable on each invoice — charged additionally on cash payments",
    "Invoices raised on milestone completion + acceptance sign-off",
    "P1 (₹6 Lakhs) due immediately on SOW signing — development continues post receipt",
    "Tech Debt terms: No interest charged. Repayment triggered when SiboniTech closes funding round or reaches profitability (as mutually agreed)",
    "Mohit Prasad retains right to convert Tech Debt to equity (at fair value) by mutual written agreement",
    "All amounts in Indian Rupees (INR). Project cost is estimated; actual invoiced amounts may vary subject to scope changes",
]
for item in pt_items:
    p2 = ptc.add_paragraph()
    set_para_spacing(p2, before=1, after=1)
    run = p2.add_run("•  " + item)
    set_run_font(run, size=9.5, color=BLACK)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 8 — AWS DEPLOYMENT
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "8.  AWS Deployment Scope & Cost Breakdown", 1, BLUE)
add_rule(doc)

add_para(doc,
    "The total project value of ₹90 Lakhs includes all AWS deployment, infrastructure setup, "
    "DevOps engineering, and 12-month initial operational costs. The following breakdown "
    "details the AWS architecture and cost components covered under Phase 4.",
    before=6, after=8, size=10)

aws_table = doc.add_table(rows=13, cols=4)
aws_table.style = 'Table Grid'
aws_table.columns[0].width = Inches(2.0)
aws_table.columns[1].width = Inches(2.5)
aws_table.columns[2].width = Inches(1.2)
aws_table.columns[3].width = Inches(0.8)

add_table_header_row(aws_table, ["AWS Service", "Purpose", "Est. Monthly", "Incl. in SOW"])
aws_services = [
    ("Amazon S3 + CloudFront", "React SPA hosting + global CDN edge delivery",          "₹ 1,200",  "✅"),
    ("ECS Fargate",            "MCP server containers — auto-scaling, no EC2 mgmt",     "₹ 8,500",  "✅"),
    ("RDS PostgreSQL (Multi-AZ)","Decision history, audit logs, user data, persistence", "₹ 12,000", "✅"),
    ("ElastiCache Redis",      "Real-time data caching — sub-second KPI responses",     "₹ 6,500",  "✅"),
    ("API Gateway + Lambda",   "Serverless connector endpoints per data source",        "₹ 4,200",  "✅"),
    ("Route 53 + ACM",         "Custom domain DNS management + SSL/TLS certificate",    "₹ 600",    "✅"),
    ("AWS WAF + Shield Std",   "Web application firewall + DDoS protection",            "₹ 3,800",  "✅"),
    ("CloudWatch",             "Monitoring, logging, alerting, custom dashboards",      "₹ 2,500",  "✅"),
    ("AWS Secrets Manager",    "API keys, credentials, MCP connector auth tokens",      "₹ 800",    "✅"),
    ("ECR (Container Registry)","Docker image storage for ECS Fargate deployments",     "₹ 400",    "✅"),
    ("VPC + NAT Gateway",      "Network isolation, private subnets, NAT for egress",    "₹ 3,200",  "✅"),
    ("Data Transfer (Egress)", "Outbound traffic to end users and connector APIs",      "₹ 2,800",  "✅"),
]
for i, (svc, purpose, monthly, incl) in enumerate(aws_services):
    row = aws_table.rows[i+1]
    shade_cell(row.cells[0], 'EFF6FF' if i%2==0 else 'FFFFFF')
    shade_cell(row.cells[3], 'F0FDF4')
    cell_para(row.cells[0], svc, size=9.5, bold=True, color=BLUE)
    cell_para(row.cells[1], purpose, size=9)
    cell_para(row.cells[2], monthly, size=9.5, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(row.cells[3], incl, size=10, bold=True, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER)

set_para_spacing(doc.add_paragraph(), before=4, after=0)

add_para(doc, "Total estimated monthly AWS infrastructure cost: ₹ 46,500 – ₹ 75,000 "
    "(depending on data volumes and API call frequency). Annual AWS cost for 12 months "
    "included in Phase 4 budget allocation of ₹19 Lakhs.",
    size=10, italic=True, color=GREY_TX, before=4, after=4)

add_para(doc, "CI/CD Pipeline:", bold=True, size=10.5, color=NAVY, before=6, after=2)
ci_items = [
    "GitHub Actions workflow → Docker build → Push to ECR → ECS Blue/Green deployment",
    "Automated tests run on every pull request (build verification)",
    "Environment separation: development, staging, production",
    "Infrastructure-as-Code using AWS CDK / CloudFormation templates",
    "Deployment rollback capability within 5 minutes",
]
for item in ci_items:
    add_bullet(doc, item, size=9.5)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 9 — ASSUMPTIONS
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "9.  Assumptions, Dependencies & Exclusions", 1, BLUE)
add_rule(doc)

add_heading(doc, "Assumptions", 3, NAVY)
assumptions = [
    "Client will provide access to enterprise systems (SAP, Salesforce, Workday) credentials and API documentation within 5 business days of Phase 2 kickoff",
    "Client has valid licences for all enterprise systems being integrated",
    "Client will provide AWS account access (or create a dedicated account) for Phase 4 deployment",
    "A dedicated point of contact from SiBoNiTech will be available for weekly reviews and UAT sign-offs",
    "All API endpoints, data schemas, and sandbox environments for enterprise systems will be made available by client before Phase 2 development begins",
    "Real-time data feed licences for Reuters and Bloomberg are the responsibility of the Client",
    "OpenAI / Azure OpenAI API credentials and cost will be managed by Client; setup included in Phase 3 scope",
]
for a in assumptions:
    add_bullet(doc, a, size=9.5)

add_heading(doc, "Dependencies", 3, NAVY)
deps = [
    "MCP Connector SDK (open-source) — no licence cost",
    "React 18 + Vite 5 — MIT licence, no cost",
    "python-pptx, python-docx — MIT licence, no cost",
    "AWS account with appropriate service limits and billing enabled",
    "Enterprise API rate limits must support real-time sync intervals (15-minute minimum)",
]
for d in deps:
    add_bullet(doc, d, size=9.5)

add_heading(doc, "Exclusions", 3, NAVY)
excls = [
    "Enterprise software licences (SAP, Salesforce, Workday, Bloomberg, Reuters) — Client responsibility",
    "Ongoing AWS monthly infrastructure costs post-project handover",
    "Ongoing OpenAI / LLM API costs",
    "Hardware procurement of any kind",
    "Physical network or data centre infrastructure",
    "ERP/CRM data migration or data cleansing",
    "Custom mobile app development (iOS / Android) — available as add-on engagement",
    "Support beyond the 90-day hypercare period without a separate maintenance contract",
]
for e in excls:
    add_bullet(doc, e, size=9.5)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 10 — ACCEPTANCE CRITERIA
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "10.  Acceptance Criteria", 1, BLUE)
add_rule(doc)

add_para(doc,
    "Each milestone is considered accepted when the following criteria are met and "
    "a formal Acceptance Certificate is signed by both parties.",
    before=6, after=6, size=10, color=GREY_TX)

criteria = {
    "Phase 1 (Current)": [
        "All 10 pages render correctly across Chrome, Edge, Firefox (desktop)",
        "All 4 user roles log in and see correct role-based navigation",
        "Insights Hub loads KPI data and drill-down cards within 2 seconds",
        "Analyst Studio generates insight from prompt and Share with CEO pushes to CEO view",
        "Admin Panel shows all 9 connectors, system prompts can be toggled and run",
        "Board Brief auto-generates with current KPI data",
        "Snap & Download functions produce correct output",
        "npm run build completes in < 10s with zero errors",
        "All code committed to GitHub main branch with clean working tree",
    ],
    "Phase 2 (MCP)": [
        "All 9 live connectors show real-time data with < 15-minute staleness",
        "MCP server uptime > 99.5% over 7-day UAT period",
        "Data schema validated against enterprise system schemas",
        "Admin Panel reflects live connector status accurately",
    ],
    "Phase 3 (AI)": [
        "Analyst Studio generates LLM-powered insights within 3 seconds",
        "System prompts execute on schedule with LLM responses stored",
        "AI narratives pass content quality review by client stakeholders",
        "Confidence scoring correlates with data source coverage",
    ],
    "Phase 4 (AWS)": [
        "Application loads in < 2 seconds globally via CloudFront CDN",
        "99.9% uptime demonstrated over 30-day observation period",
        "All data encrypted at rest (RDS) and in transit (TLS 1.3)",
        "CI/CD pipeline deploys new version within 10 minutes of merge",
        "CloudWatch dashboards operational with P1 alerting configured",
    ],
}

for phase, items in criteria.items():
    add_para(doc, phase, bold=True, size=10.5, color=NAVY, before=6, after=2)
    for item in items:
        add_bullet(doc, item, size=9.5)

# ════════════════════════════════════════════════════════════════════════════
# SECTIONS 11–14
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "11.  Intellectual Property & Maintenance", 1, BLUE)
add_rule(doc)
add_para(doc,
    "Upon full settlement of all cash payments (₹39,00,000) and formal documentation of the "
    "Tech Debt agreement (₹51,00,000), full intellectual property rights — including source "
    "code, design assets, data models, documentation, and all related artefacts — shall "
    "transfer entirely to SiBoNiTech Pvt. Ltd.",
    before=6, after=4)
add_para(doc,
    "Prior to full cash payment, Mohit Prasad retains all IP rights. The Client is granted "
    "a limited, non-transferable licence to use the delivered software for internal evaluation, "
    "UAT, and production purposes only. The Client may not sub-licence, resell, or "
    "open-source any deliverable without written consent from Mohit Prasad.",
    before=4, after=4)
add_para(doc,
    "Phase-wise IP transfer: Upon payment of each milestone's cash component, the Client gains "
    "irrevocable licence to the deliverables of that specific milestone for internal production use.",
    before=4, after=4, italic=True, color=GREY_TX)
add_para(doc, "Maintenance & Future Upgrades:", bold=True, color=NAVY, size=10.5, before=8, after=2)
add_para(doc,
    "All future maintenance, enhancements, upgrades, and new feature development shall be "
    "handled exclusively by Mohit Prasad, either as an individual or through his OPC / company "
    "(details to be furnished). SiboniTech agrees to engage Mohit Prasad as the preferred "
    "technology partner for all platform-related work for a minimum period of three (3) years "
    "post go-live, unless mutually terminated by written agreement.",
    before=2, after=4)
add_para(doc,
    "Maintenance Engagement Model: Time & material / retainer — to be agreed in a separate "
    "Maintenance & Support Agreement (MSA) post Phase 5 delivery. Estimated ₹2–3 Lakhs/month "
    "for ongoing support, monitoring, and upgrades.",
    before=2, after=4, italic=True, color=GREY_TX)

add_heading(doc, "12.  Confidentiality", 1, BLUE)
add_rule(doc)
add_para(doc,
    "Both parties agree to maintain strict confidentiality of all proprietary information, "
    "trade secrets, business data, and technical specifications shared under this agreement. "
    "This obligation survives termination of the SOW for a period of five (5) years.",
    before=6, after=4)
add_para(doc,
    "The Service Provider shall not disclose the existence of this engagement, client name, "
    "or any project details to third parties without prior written consent, except as required "
    "by law or regulation.",
    before=4, after=4)

add_heading(doc, "13.  Termination & Dispute Resolution", 1, BLUE)
add_rule(doc)
add_para(doc, "Termination for Convenience:", bold=True, color=NAVY, before=6, after=2)
add_para(doc,
    "Either party may terminate this agreement with 30 days' written notice. In such event, "
    "the Client shall pay for all work completed up to the termination date at the pro-rata "
    "milestone rate. No refund of payments already made shall be due.",
    before=2, after=4)
add_para(doc, "Dispute Resolution:", bold=True, color=NAVY, before=4, after=2)
add_para(doc,
    "Any disputes shall first be escalated to senior management of both parties for resolution "
    "within 15 business days. If unresolved, disputes shall be referred to binding arbitration "
    "under the Arbitration and Conciliation Act, 1996 (India). Jurisdiction: Mumbai, Maharashtra.",
    before=2, after=6)

add_heading(doc, "14.  Signatories", 1, BLUE)
add_rule(doc)
add_para(doc,
    "By signing below, both parties agree to be bound by the terms and conditions of this "
    "Statement of Work.",
    before=6, after=10, size=10, color=GREY_TX)

sig_table = doc.add_table(rows=1, cols=2)
sig_table.style = 'Table Grid'
sig_table.columns[0].width = Inches(3.0)
sig_table.columns[1].width = Inches(3.5)

shade_cell(sig_table.rows[0].cells[0], 'EFF6FF')
shade_cell(sig_table.rows[0].cells[1], 'F0FDF4')

sig_left = sig_table.rows[0].cells[0]
sig_right = sig_table.rows[0].cells[1]

for label, cell in [("FOR: SiBoNiTech Pvt. Ltd. (CLIENT)", sig_left),
                    ("FOR: Mohit Prasad / OPC (SERVICE PROVIDER)", sig_right)]:
    p = cell.add_paragraph()
    set_para_spacing(p, before=4, after=4)
    run = p.add_run(label)
    set_run_font(run, size=9.5, bold=True, color=BLUE)
    for field in ["Name:   ___________________________________",
                  "Title:    ___________________________________",
                  "Date:   ___________________________________",
                  "Sign:   ___________________________________",
                  "Stamp: ___________________________________"]:
        fp = cell.add_paragraph()
        set_para_spacing(fp, before=8, after=2)
        run = fp.add_run(field)
        set_run_font(run, size=10, color=BLACK)

# ════════════════════════════════════════════════════════════════════════════
# FOOTER NOTE
# ════════════════════════════════════════════════════════════════════════════
set_para_spacing(doc.add_paragraph(), before=16, after=0)
add_rule(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p, before=4, after=0)
r = p.add_run(
    "SiBoNi CXO Cockpit  ·  Statement of Work v1.1  ·  Dated 29 March 2026  ·  "
    "CONFIDENTIAL  ·  Estimated Cost ₹90L  |  SiboniTech Payable ₹39L  |  Tech Debt ₹51L @ 0%"
)
set_run_font(r, size=8, color=GREY_TX)

# ════════════════════════════════════════════════════════════════════════════
out = "SiBoNi_SOW_v1.1.docx"
doc.save(out)
print(f"Saved: {out}")
